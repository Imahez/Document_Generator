import openai
import streamlit as st
from docx import Document
from fpdf import FPDF
import base64

# OpenAI API key
openai.api_key = 'your_openai_api_key'

# Function to generate content with ChatGPT
def generate_content(topic, style="formal", sections=None):
    if sections is None:
        sections = ["Introduction", "Main Content", "Conclusion"]

    prompt = f"Generate a {style} document on the topic '{topic}' with sections: {', '.join(sections)}."

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful document generator."},
            {"role": "user", "content": prompt}
        ]
    )

    content = response.choices[0].message['content']
    return content

# Function to save content as DOCX
def save_to_docx(content):
    doc = Document()
    doc.add_heading("Generated Document", level=1)
    doc.add_paragraph(content)
    
    # Save the document to a binary object
    doc_binary = doc_to_binary(doc)
    return doc_binary

def doc_to_binary(doc):
    binary_io = BytesIO()
    doc.save(binary_io)
    binary_io.seek(0)
    return binary_io.getvalue()

# Function to save content as PDF
def save_to_pdf(content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, content)
    
    pdf_binary = pdf_to_binary(pdf)
    return pdf_binary

def pdf_to_binary(pdf):
    binary_io = BytesIO()
    pdf.output(binary_io)
    binary_io.seek(0)
    return binary_io.getvalue()

# Function to create a download link
def create_download_link(data, filename, file_format):
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download {file_format}</a>'
    return href

# Streamlit app layout
st.title("Document Generator with ChatGPT")

topic = st.text_input("Enter the document topic", "The Impact of Climate Change")
style = st.selectbox("Choose the document style", ["formal", "informative", "analytical", "persuasive"])
sections = st.multiselect("Select sections for the document", ["Introduction", "Background", "Analysis", "Conclusion", "Recommendations"], default=["Introduction", "Conclusion"])
output_format = st.selectbox("Choose output format", ["docx", "pdf"])

if st.button("Generate Document"):
    with st.spinner("Generating document..."):
        content = generate_content(topic, style, sections)
        st.success("Document generated successfully!")

        st.subheader("Document Preview")
        st.write(content)

        if output_format == "docx":
            doc_binary = save_to_docx(content)
            st.markdown(create_download_link(doc_binary, "generated_document.docx", "DOCX"), unsafe_allow_html=True)
        elif output_format == "pdf":
            pdf_binary = save_to_pdf(content)
            st.markdown(create_download_link(pdf_binary, "generated_document.pdf", "PDF"), unsafe_allow_html=True)
